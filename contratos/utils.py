import pandas as pd
from docx import Document
import os
from django.conf import settings
from datetime import datetime
import re
from rapidfuzz import process, fuzz
from docx.shared import Pt

ruta_excel = settings.EXCEL_FILE_PATH
ruta_docs = settings.MEDIA_PATH

MESES_ES = {
    "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
    "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
    "September": "septiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
}

# ✅ FORMATO DE FECHAS Y MONEDAS
def formatear_fecha(fecha):
    if pd.isna(fecha) or fecha == "":
        return "N/A"
    try:
        fecha_obj = pd.to_datetime(fecha)
        dia = fecha_obj.day
        mes = MESES_ES[fecha_obj.strftime("%B")]
        año = fecha_obj.strftime("%y")
        return f"{dia:02d}-{mes}-{año}"
    except Exception:
        return "N/A"

def obtener_fecha_hoy():
    hoy = datetime.today()
    dia = hoy.day
    mes = MESES_ES[hoy.strftime("%B")]
    año = hoy.year
    return f"Ciudad de México, {dia} de {mes} de {año}"

def formatear_moneda(valor):
    try:
        valor = safe_float(valor, 0)
        return f"${float(valor):,.2f}" if valor not in ["", "N/A"] else "N/A"
    except Exception as e:
        print(f"❗ Error al formatear moneda: {e}")
        return "N/A"

# ✅ LIMPIEZA DE NOMBRES Y RFC
def limpiar_nombre_proveedor(nombre):
    if not isinstance(nombre, str):
        return nombre
    nombre = re.sub(r'\b(S\.?A\.? DE C\.?V\.?|S\.? DE R\.?L\.?|S\.?C\.?|SAPI DE C\.?V\.?)\b', '', nombre, flags=re.IGNORECASE)
    return " ".join(nombre.split()).strip()

def corregir_nombre_proveedor(nombre, lista_proveedores):
    mejor_coincidencia, puntuacion, _ = process.extractOne(nombre, lista_proveedores, scorer=fuzz.token_sort_ratio)
    return mejor_coincidencia if puntuacion >= 80 else nombre

def buscar_rfc(nombre_proveedor, df_rfc):
    if "PROVEEDOR" in df_rfc.columns:
        df_rfc["PROVEEDOR"] = df_rfc["PROVEEDOR"].map(lambda x: limpiar_nombre_proveedor(x) if isinstance(x, str) else x)
        lista_proveedores = df_rfc["PROVEEDOR"].dropna().unique().tolist()
        nombre_proveedor_limpio = limpiar_nombre_proveedor(nombre_proveedor)
        nombre_proveedor_corregido = corregir_nombre_proveedor(nombre_proveedor_limpio, lista_proveedores)
        rfc_resultado = df_rfc[df_rfc["PROVEEDOR"].str.lower() == nombre_proveedor_corregido.lower()]
        return rfc_resultado.iloc[0]["RFC"] if not rfc_resultado.empty else "N/A"
    return "Error en datos"

# ✅ DESTINATARIOS
def obtener_destinatarios():
    try:
        df = pd.read_excel(ruta_excel, sheet_name="Destinatarios").fillna("")
        for d in df.to_dict(orient="records"):
            d["NOMBRE"] = d["NOMBRE"].strip().replace("\u00A0", " ")
        return df.to_dict(orient="records")
    except Exception:
        return []


def safe_float(value, default=0):
    try:
        if value is None or value == "":
            return default
        return float(value)
    except Exception as e:
        print(f"❗ Error en safe_float con valor {value}: {e}")
        return default


# ✅ BÚSQUEDAS
def buscar_contrato_en_excel(contrato_id):
    año_contrato = extraer_año_contrato(contrato_id)
    print("📅 Año detectado:", año_contrato)
    if not año_contrato:
        print("❌ Año no detectado.")
        return None

    try:
        hojas_excel = pd.ExcelFile(ruta_excel).sheet_names
        print("📋 Hojas disponibles:", hojas_excel)
        if año_contrato not in hojas_excel:
            print(f"❌ Hoja '{año_contrato}' no encontrada.")
            return None

        df = pd.read_excel(ruta_excel, sheet_name=año_contrato).fillna("")
        print(f"📑 {año_contrato}: {len(df)} registros cargados")
        resultado = df[df["CONTRATO"].str.strip() == contrato_id.strip()]
        print("🔍 Filtrado por contrato:", resultado)

        if resultado.empty:
            print("❌ No se encontró el contrato en el DataFrame.")
            return None

        row = resultado.iloc[0]
        resultado_dict = row.to_dict()
        
        print("📝 Resultado dict inicial:", resultado_dict)
        
        try:

            resultado_dict.update({
                "IMPORTE_MAX_SIN_IVA": formatear_moneda(row.get("IMPORTE MAXIMO (SIN IVA)")),
                "IVA": formatear_moneda(row.get("IVA")),
                "TOTAL_MAX": formatear_moneda(row.get("TOTAL MAXIMO")),
                "IMPORTE_MIN_SIN_IVA": formatear_moneda(row.get("IMPORTE MÍNIMO (SIN IVA)")),
                "IVA2": formatear_moneda(row.get("IVA2")),
                "TOTAL_MIN": formatear_moneda(row.get("TOTAL MINIMO")),
                "FECHA_INICIO": formatear_fecha(row.get("INICIO DE VIGENCIA")),
                "FECHA_FIN": formatear_fecha(row.get("FIN DE VIGENCIA")),
                "OBSERVACIONES": row.get("OBSERVACIONES", "N/A"),
                "FIRMA_PROVEEDOR": formatear_fecha(row.get("FIRMA PROVEEDOR")),
                "ENVIO_FIRMA_ADMIN": formatear_fecha(row.get("SE ENVIA A FIRMA DEL ADMINISTRADOR")),
                "DEVUELTO_ADMIN": formatear_fecha(row.get("ENVÍAN DEL ADMINISTRADOR EL CONTRATO")),
                "ENVIO_DG": formatear_fecha(row.get("SE ENVIA A FIRMA DEL DIRECTOR GENERAL")),
                "DEVUELTO_DG": formatear_fecha(row.get("ENVÍAN DE D.G. EL CONTRATO")),
                "FECHA_FORMALIZACION_TESO": formatear_fecha(row.get("FECHA DE FORMALIZACIÓN DE FIANZA EN TESORERÍA y/o CONTRATO")),
                "AREA_SOLICITANTE": row.get("AREA SOLICITANTE", "N/A"),
            })
    
            try:
                total_maximo = safe_float(row.get("TOTAL MAXIMO", 0))
            except Exception as e:
                print("❗ Error en total_maximo:", e)
                total_maximo = 0
            
            
            
            resultado_dict.update({
                "FIANZA_DE_CUMPLIMIENTO": calcular_monto(row.get("FIANZA DE CUMPLIMIENTO"), total_maximo),
                "POLIZA_DE_RESPONSABILIDAD_CIVIL": calcular_monto(row.get("POLIZA DE RESPONSABILIDAD CIVIL"), total_maximo),
                "VICIOS_OCULTOS": calcular_monto(row.get("VICIOS OCULTOS"), total_maximo),
                "FIANZA_DE_ANTICIPO": calcular_monto(row.get("FIANZA DE ANTICIPO"), total_maximo),
                "REPSE": calcular_monto(row.get("REPSE"), total_maximo)
            })
        except Exception as e:
            print(f"❗ Error al actualizar resultado_dict: {e}")
            return None

        try:
            df_rfc = pd.read_excel(ruta_excel, sheet_name="RFC")
            resultado_dict["RFC"] = buscar_rfc(resultado_dict["PROVEEDOR"], df_rfc)
        except Exception:
            resultado_dict["RFC"] = "Error RFC"

        if resultado_dict.get("PLURI", "").strip().upper() == "SI":
            try:
                df_pluri = pd.read_excel(ruta_excel, sheet_name="Plurianuales").fillna("")
                pluri_resultado = df_pluri[df_pluri["CONTRATO"] == contrato_id]
                if not pluri_resultado.empty:
                    resultado_dict["PLURIANUAL_INFO"] = pluri_resultado.iloc[0].to_dict()
            except Exception:
                pass
            
        return resultado_dict
    except Exception as e:
        print(f"❗ Error general en buscar_contrato_en_excel: {e}")
        return None

def buscar_contratos_por_proveedor(proveedor, anio=None):
    contratos = []
    xl = pd.ExcelFile(ruta_excel)

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name).fillna("")
        # Verificar que exista la columna 'PROVEEDOR'
        if "PROVEEDOR" not in df.columns or "CONTRATO" not in df.columns:
            continue  # saltar la hoja que no tiene la estructura correcta
        df_filtrado = df[df["PROVEEDOR"].str.contains(proveedor, case=False, na=False)]
        if anio:
            df_filtrado = df_filtrado[df_filtrado["CONTRATO"].str.contains(f"/{anio}")]
        for _, row in df_filtrado.iterrows():
            contratos.append({
                "CONTRATO": row["CONTRATO"],
                "PROVEEDOR": row["PROVEEDOR"],
                "DESCRIPCIÓN": row.get("DESCRIPCIÓN", "N/A")
            })

    return contratos

def buscar_convenios(proveedor, anio=None):
    convenios = []
    xl = pd.ExcelFile(ruta_excel)

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name).fillna("")
        df.columns = df.columns.str.strip().str.upper()
        if "PROVEEDOR" not in df.columns or "CONTRATO" not in df.columns:
            # Saltamos hojas que no tengan la columna correcta
            print(f"❌ La hoja '{sheet_name}' NO contiene la columna 'PROVEEDOR'")
            continue
        
        df_filtrado = df[
            (df["PROVEEDOR"].str.contains(proveedor, case=False, na=False)) &
            (df["CONTRATO"].str.contains(r'-[IVXLCDM]+/', na=False))
        ]
        if anio:
            #df_filtrado = df_filtrado[df_filtrado["CONTRATO"].str.endswith(f"/{str(anio)[-2:]}")]
            df_filtrado = df_filtrado[df_filtrado["CONTRATO"].str.contains(f"/{anio[-2:]}", na=False)]
        for _, row in df_filtrado.iterrows():
            convenios.append({
                "CONTRATO": row["CONTRATO"],
                "PROVEEDOR": row["PROVEEDOR"],
                "DESCRIPCIÓN": row.get("DESCRIPCIÓN", "N/A")
            })

    return convenios

def extraer_año_contrato(contrato_id):
    partes = contrato_id.split("/")
    for parte in partes:
        if parte.isdigit() and parte in ["2020", "2021", "2022", "2023", "2024", "2025"]:
            return parte
    return None

def calcular_monto(valor, total_maximo):
    try:
        if pd.isna(valor) or valor == "" or total_maximo == 0:
            return "NO APLICA"
        if isinstance(valor, str) and "%" in valor:
            valor = valor.replace("%", "").strip()
        monto = (float(valor)) * float(total_maximo)
        return f"${monto:,.2f}"
    except Exception:
        return "NO APLICA"

# ✅ GENERAR DOCUMENTOS
def generar_documento(tipo, datos_contrato, destinatario, ccp_lista=None):
    output_dir = os.path.join(settings.MEDIA_ROOT, "DocsGenerados")
    os.makedirs(output_dir, exist_ok=True)

    plantilla_path = os.path.join(settings.BASE_DIR, "contratos/doc_templates", f"{tipo}.docx")
    output_filename = f"{tipo}_{datos_contrato['CONTRATO'].replace('/', '_')}.docx"
    output_path = os.path.join(output_dir, output_filename)

    if not os.path.exists(plantilla_path):
        print(f" Plantilla no encontrada: {plantilla_path}")
        return None

    doc = Document(plantilla_path)
    fecha_hoy = obtener_fecha_hoy()

    reemplazos = {
        "{CONTRATO}": datos_contrato.get("CONTRATO", "N/A"),
        "{PROVEEDOR}": datos_contrato.get("PROVEEDOR", "N/A"),
        "{DESCRIPCION}": datos_contrato.get("DESCRIPCIÓN", "N/A"),
        "{FECHA_HOY}": fecha_hoy,
        "{NOMBRE}": destinatario.get("NOMBRE", "N/A"),
        "{CARGO}": destinatario.get("CARGO", "N/A"),
    }

    if tipo == "poliza":
        reemplazos["{INICIO_VIGENCIA}"] = datos_contrato.get("FECHA_INICIO", "N/A")
        reemplazos["{FIN_VIGENCIA}"] = datos_contrato.get("FECHA_FIN", "N/A")

    # C.C.P
    ccp_texto = "\n".join([
        f"{entry.split('||')[0].strip()}. {entry.split('||')[1].strip()}." for entry in ccp_lista]) if ccp_lista else ""

    for p in doc.paragraphs:
        for key, value in reemplazos.items():
            p.text = p.text.replace(key, str(value))
        if "{C.C.P}" in p.text:
            p.text = ""
            par = p
            run_label = par.add_run("C.c.p.-")
            run_label.font.size = Pt(7)
            run_label.font.name = "Calibri"
            run_label.bold = False

            if ccp_lista:
                for idx, entry in enumerate(ccp_lista):
                    nombre, cargo = entry.split("||")
                    newline = doc.add_paragraph()
                    newline.paragraph_format.space_after = Pt(0)
                    run_entry = newline.add_run(f"{nombre.title().strip()}. {cargo.title().strip()}.")
                    run_entry.font.size = Pt(7)
                    run_entry.font.name = "Calibri"
                    run_entry.bold = True

    doc.save(output_path)
    return f"{settings.MEDIA_URL}DocsGenerados/{output_filename}"

def buscar_pedido_en_excel(numero_pedido):
    try:
        xl = pd.ExcelFile(settings.EXCEL_PEDIDOS_PATH)
        if "PEDIDOS" not in xl.sheet_names:
            print("❌ La hoja 'PEDIDOS' no está en el archivo.")
            return []

        df = xl.parse("PEDIDOS").fillna("")
        df.columns = df.columns.str.strip().str.upper()
        resultado = df[df["N° PEDIDO"].astype(str).str.strip() == numero_pedido]
        if resultado.empty:
            return []
        return resultado[[
            "N° SERVICIO",
            "PROVEEDOR",
            "ÁREA SOLICITANTE",
            "DESCRIPCIÓN",
            "FECHA DE LA ORDEN",
            "FIRMA PROVEED.",
            "SE ENVIA A FIRMA DEL ADMIN.",
            "REGRESAN CON FIRMA DEL ADMIN.",
            "FECHA FORMALIZACIÓN (REAL)"
        ]].to_dict(orient="records")

    except Exception as e:
        print(f"❗ Error al buscar pedido: {e}")
        return []

def buscar_orden_en_excel(numero_servicio):
    try:
        xl = pd.ExcelFile(settings.EXCEL_PEDIDOS_PATH)
        if "SERVICIOS" not in xl.sheet_names:
            print("❌ La hoja 'SERVICIOS' no está en el archivo.")
            return []

        df = xl.parse("SERVICIOS").fillna("")
        df.columns = df.columns.str.strip().str.upper()

        resultado = df[df["N° SERVICIO"].astype(str).str.strip() == numero_servicio]
        if resultado.empty:
            return []

        return resultado[[
            "N° SERVICIO",
            "PROVEEDOR",
            "ÁREA SOLICITANTE",
            "DESCRIPCIÓN",
            "FECHA DE LA ORDEN",
            "FIRMA PROVEED.",
            "SE ENVIA A FIRMA DEL ADMIN.",
            "REGRESAN CON FIRMA DEL ADMIN.",
            "FECHA FORMALIZACIÓN (REAL)"
        ]].to_dict(orient="records")

    except Exception as e:
        print(f"❗ Error al buscar orden: {e}")
        return []